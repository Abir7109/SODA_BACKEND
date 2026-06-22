import os, csv, json, io, asyncio
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EXPORT_DIR = Path.home() / "Desktop" / "soda_exports"

_SODA_BG = RGBColor(0x0A, 0x0E, 0x17)
_SODA_CYAN = RGBColor(0x00, 0xFB, 0xFB)
_SODA_TEXT = RGBColor(0xCC, 0xCC, 0xCC)
_SODA_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_SODA_LINE = RGBColor(0x1A, 0x2A, 0x3A)

def _ensure_export_dir():
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

def _data_to_rows(data):
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except json.JSONDecodeError:
            return [{"content": data}]
    if isinstance(data, dict):
        if all(isinstance(v, (list, dict, str, int, float, bool)) for v in data.values()):
            return [data]
        for k, v in data.items():
            if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                return v
        return [data]
    if isinstance(data, list):
        if len(data) == 0:
            return [{"content": "No data"}]
        if all(isinstance(x, dict) for x in data):
            return data
        return [{"item": str(x)} for x in data]
    return [{"data": str(data)}]

def _write_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def _set_cell_text(cell, text, font_name="Consolas", font_size=9, bold=False, color=_SODA_TEXT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    _remove_space_before(p)
    _remove_space_after(p)

def _remove_space_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")

def _remove_space_after(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:after"), "40")

def _set_page_background(doc, color_hex):
    sectPr = doc.sections[0]._sectPr
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), color_hex)
    sectPr.append(bg)

def _add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 18, 2: 14, 3: 12}
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = True
    run.font.color.rgb = _SODA_CYAN
    _remove_space_after(p)
    return p

def _add_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "00fbfb")
    pBdr.append(bottom)
    pPr.append(pBdr)
    _remove_space_before(p)
    _remove_space_after(p)

def _add_body_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = _SODA_TEXT
    _remove_space_after(p)

def _resolve_path(path, fallback_dir, title, ext):
    if path:
        p = Path(str(path).replace("~", str(Path.home())))
        if p.suffix:
            p.parent.mkdir(parents=True, exist_ok=True)
            return p
        p.mkdir(parents=True, exist_ok=True)
        return p / f"{title}{ext}"
    fallback_dir.mkdir(parents=True, exist_ok=True)
    return fallback_dir / f"{title}{ext}"

def export_json(data, title, path=None):
    rows = _data_to_rows(data)
    out = _resolve_path(path, EXPORT_DIR, title, ".json")
    out.write_text(json.dumps(rows if len(rows) > 1 else data, indent=2, ensure_ascii=False, default=str), encoding="utf-8")
    return {"success": True, "path": str(out), "message": f"Saved JSON with {len(rows)} records"}

def export_html(data, title, path=None):
    rows = _data_to_rows(data)
    scores = data.get("scores", {}) if isinstance(data, dict) else {}
    vitals = data.get("vitals", {}) if isinstance(data, dict) else {}
    opportunities = data.get("opportunities", []) if isinstance(data, dict) else []
    passed_audits = data.get("passed_audits", []) if isinstance(data, dict) else []
    url = data.get("url", "") if isinstance(data, dict) else ""
    strategy = data.get("strategy", "desktop") if isinstance(data, dict) else ""
    scoreColor = lambda v: "#00ff88" if v >= 90 else "#ffaa00" if v >= 50 else "#ff3355"
    parts = ['<!DOCTYPE html><html><head><meta charset="utf-8"><title>PageSpeed Report</title>']
    parts.append("<style>body{font-family:monospace;background:#0a0a0f;color:#e0e0e0;padding:20px;max-width:800px;margin:0 auto}")
    parts.append("h1{color:#00f0ff;border-bottom:1px solid #1e1e2e;padding-bottom:8px}")
    parts.append("h2{color:#00f0ff;margin-top:24px}")
    parts.append("table{width:100%;border-collapse:collapse;margin:8px 0}")
    parts.append("td,th{border:1px solid #1e1e2e;padding:6px 10px;text-align:left;font-size:13px}")
    parts.append("th{color:#666680;text-transform:uppercase;font-size:11px}")
    parts.append(".pass{color:#00ff88}.fail{color:#ff3355}.warn{color:#ffaa00}")
    parts.append("</style></head><body>")
    if url:
        parts.append(f"<h1>PageSpeed Insights Report</h1>")
        parts.append(f'<p style="color:#666680">{url} | {strategy}</p>')
        if scores:
            parts.append("<h2>Scores</h2><table>")
            for k, v in scores.items():
                c = scoreColor(v)
                parts.append(f'<tr><td>{k.replace("_"," ").upper()}</td><td><span style="background:{c}20;color:{c};padding:2px 8px;font-weight:700">{v}%</span></td></tr>')
            parts.append("</table>")
        if vitals:
            parts.append("<h2>Core Web Vitals</h2><table>")
            for k, v in vitals.items():
                parts.append(f"<tr><td>{k.upper()}</td><td>{v}</td></tr>")
            parts.append("</table>")
        if opportunities:
            parts.append(f"<h2>Optimization Opportunities ({len(opportunities)})</h2>")
            for o in opportunities:
                parts.append(f'<div style="border:1px solid #1e1e2e;padding:10px;margin:6px 0;background:#0a0a0f">')
                parts.append(f'<span style="font-weight:700;color:#fff">{o.get("title","")}</span> <span style="color:#ff3355;float:right">-{o.get("score_impact",0)}pt</span>')
                if o.get("issue"):
                    parts.append(f'<br><span style="color:#b0b0cc;font-size:12px">{o["issue"]}</span>')
                parts.append("</div>")
        if passed_audits:
            parts.append(f"<h2>Passed Audits ({len(passed_audits)})</h2><ul>")
            for t in passed_audits:
                parts.append(f'<li style="color:#00ff88">{t}</li>')
            parts.append("</ul>")
    else:
        parts.append(f"<h1>{title}</h1>")
        if len(rows) == 1 and list(rows[0].keys()) == ["content"]:
            parts.append(f"<pre>{rows[0]['content']}</pre>")
        else:
            headers = list(rows[0].keys())
            parts.append("<table><tr>" + "".join(f"<th>{h.replace('_',' ').title()}</th>" for h in headers) + "</tr>")
            for row in rows:
                parts.append("<tr>" + "".join(f"<td>{str(row.get(h,''))}</td>" for h in headers) + "</tr>")
            parts.append("</table>")
    parts.append("</body></html>")
    out = _resolve_path(path, EXPORT_DIR, title, ".html")
    out.write_text("\n".join(parts), encoding="utf-8")
    return {"success": True, "path": str(out), "message": f"Saved HTML with {len(rows)} records"}

def export_markdown(data, title, path=None):
    rows = _data_to_rows(data)
    lines = [f"# {title}", "", "---", ""]
    if len(rows) == 1 and list(rows[0].keys()) == ["content"]:
        lines.append(rows[0]["content"])
        lines.append("")
    else:
        headers = list(rows[0].keys())
        lines.append("| " + " | ".join(h.replace("_", " ").title() for h in headers) + " |")
        lines.append("| " + " | ".join("---" for _ in headers) + " |")
        for row in rows:
            vals = [str(row.get(h, "")) for h in headers]
            lines.append("| " + " | ".join(v.replace("\n", " ")[:80] for v in vals) + " |")
        lines.append("")
        lines.append(f"*{len(rows)} records*")
        lines.append("")
    out = _resolve_path(path, EXPORT_DIR, title, ".md")
    out.write_text("\n".join(lines), encoding="utf-8")
    return {"success": True, "path": str(out), "message": f"Saved markdown with {len(rows)} records"}

def export_csv(data, title, path=None):
    rows = _data_to_rows(data)
    out = _resolve_path(path, EXPORT_DIR, title, ".csv")
    if not rows:
        out.write_text("", encoding="utf-8")
        return {"success": True, "path": str(out), "message": "Empty CSV"}
    headers = list(rows[0].keys())
    with open(out, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=headers)
        w.writeheader()
        w.writerows({h: row.get(h, "") for h in headers} for row in rows)
    return {"success": True, "path": str(out), "message": f"Saved CSV with {len(rows)} records"}

def export_docx(data, title, path=None):
    rows = _data_to_rows(data)
    out = _resolve_path(path, EXPORT_DIR, title, ".docx")
    doc = Document()
    _set_page_background(doc, "0A0E17")
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    _add_heading_custom(doc, title, 1)
    _add_line(doc)
    if len(rows) == 1 and list(rows[0].keys()) == ["content"]:
        _add_body_text(doc, rows[0]["content"])
    else:
        headers = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            _write_cell_shading(cell, "001a2e")
            _set_cell_text(cell, h.replace("_", " ").title(), bold=True, color=_SODA_CYAN, font_size=9)
            width = Inches(9.0 / len(headers))
            cell.width = width
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, h in enumerate(headers):
                cell = row_cells[i]
                val = str(row_data.get(h, ""))
                _set_cell_text(cell, val, font_size=8)
        _add_line(doc)
        p = doc.add_paragraph()
        run = p.add_run(f"{len(rows)} records  |  SODA Export")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = _SODA_LINE
    doc.save(str(out))
    return {"success": True, "path": str(out), "message": f"Saved DOCX with {len(rows)} records"}

async def export_data(data, export_format, title, path=None):
    if not title:
        title = "soda_export"
    title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)[:60]
    export_format = export_format.lower().strip()
    if export_format in ("json",):
        return export_json(data, title, path)
    elif export_format in ("html", "htm"):
        return export_html(data, title, path)
    elif export_format in ("md", "markdown", "mdown"):
        return export_markdown(data, title, path)
    elif export_format in ("csv",):
        return export_csv(data, title, path)
    elif export_format in ("docx", "doc", "word"):
        result = await asyncio.to_thread(export_docx, data, title, path)
        return result
    else:
        return {"success": False, "message": f"Unsupported format: {export_format}. Use json, html, markdown, csv, or docx."}
